import xlrd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

src = "./提案一览表.xls"
export_dir = "./dst/"

#设置表格的边框
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

data = xlrd.open_workbook(src)
xls_table = data.sheets()[0]
lines = xls_table.nrows

for i in range(1,lines):
    case_number = xls_table.cell_value(i,0)
    case_content = xls_table.cell_value(i,2)
    execute_department = xls_table.cell_value(i,3)
    reply_content = xls_table.cell_value(i,5)
    file_name = export_dir + case_number+"号 "+case_content + ".docx"
    print("Processing " + file_name)

    doc = Document()

    # 增加标题表格
    table = doc.add_table(rows=1, cols=1)

    # 表格整体居中
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 选中第一个单元格（行，列）
    title_cell = table.cell(0,0)

    # 设置表格仅有下边框,红色
    set_cell_border(
            title_cell,
            bottom={"sz": 10, "val": "single", "color": "#FF0000", "space": "0"},
        )
    
    # 增加标题内容
    t = title_cell.paragraphs[0]
    t.text = execute_department

    # 将表格内容居中
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 更改标题文字格式字体
    t.style = "Heading 1"
    t.style.font.size = Pt(28)
    t.style.font.bold = True
    t.style.font.name = u'宋体'
    t.style._element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
    t.style.font.color.rgb = RGBColor(255,0,0)

    # 增加正文内容并修改字体
    p = doc.add_paragraph(reply_content)
    p.style.font.size = Pt(14)
    p.style.font.name = u'仿宋_GB2312'
    p.style.font.color.rgb = RGBColor(0,0,0)
    p.style._element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋_GB2312')

    doc.save(file_name)